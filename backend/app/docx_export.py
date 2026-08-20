"""Word 导出（python-docx）：Markdown 风格的文书内容 → .docx（中文宋体）。"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_east_asia(style_or_run, font_name: str = "宋体") -> None:
    try:
        style_or_run.font.name = font_name
        rpr = style_or_run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def export_docx(title: str, content: str) -> BytesIO:
    doc = Document()
    _set_east_asia(doc.styles["Normal"], "宋体")
    doc.styles["Normal"].font.size = Pt(12)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_east_asia(h.runs[0], "黑体") if h.runs else None

    for line in (content or "").splitlines():
        line = line.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            _set_east_asia(p.runs[0], "黑体") if p.runs else None
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            _set_east_asia(p.runs[0], "黑体") if p.runs else None
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            _set_east_asia(p.runs[0], "黑体") if p.runs else None
        else:
            para = doc.add_paragraph(stripped)
            for run in para.runs:
                _set_east_asia(run, "宋体")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _append_markdown(doc: Document, content: str) -> None:
    """把 Markdown 风格内容追加到文档（供分节复用）。"""
    for line in (content or "").splitlines():
        line = line.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            _set_east_asia(p.runs[0], "黑体") if p.runs else None
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            _set_east_asia(p.runs[0], "黑体") if p.runs else None
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            _set_east_asia(p.runs[0], "黑体") if p.runs else None
        else:
            para = doc.add_paragraph(stripped)
            for run in para.runs:
                _set_east_asia(run, "宋体")


def export_all_docx(patient: dict, visit: dict, docs: list[dict],
                    type_names: dict | None = None) -> BytesIO:
    """完整病历导出：封面（患者+住院信息）+ 各文书分节（分页符分隔）。"""
    type_names = type_names or {}
    doc = Document()
    _set_east_asia(doc.styles["Normal"], "宋体")
    doc.styles["Normal"].font.size = Pt(12)

    # 封面
    title = doc.add_heading("完整病历", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_east_asia(title.runs[0], "黑体") if title.runs else None
    for line in [
        f"患者：{patient.get('姓名','') or patient.get('脱敏编号','')}（{patient.get('性别','')}，{patient.get('年龄','')}岁）",
        f"病案号：{patient.get('病案号','') or '—'}",
        f"住院号：{visit.get('住院号','')}",
        f"入院日期：{visit.get('入院日期','')}  出院日期：{visit.get('出院日期','') or '—'}",
        f"入院诊断：{visit.get('入院诊断','') or '—'}",
        f"出院诊断：{visit.get('出院诊断','') or '—'}",
        f"过敏史：{patient.get('过敏史','') or '—'}",
    ]:
        p = doc.add_paragraph(line)
        for run in p.runs:
            _set_east_asia(run, "宋体")

    if not docs:
        p = doc.add_paragraph("（暂无已确认文书）")
        for run in p.runs:
            _set_east_asia(run, "宋体")

    for d in docs:
        doc.add_page_break()
        name = type_names.get(d.get("doc_type"), d.get("doc_type"))
        h = doc.add_heading(f"{name}（{d.get('doc_date','')}）", level=1)
        _set_east_asia(h.runs[0], "黑体") if h.runs else None
        _append_markdown(doc, d.get("content", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
